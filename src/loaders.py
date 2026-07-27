"""Belge dosyalarını okunabilir metne çevirir.

Her sayfanın numarasını koruyoruz, çünkü cevabın altında "[dosya, s.12]" diye
kaynak göstereceğiz. Sayfa numarası olmadan kaynak gösterimi işe yaramaz.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path


@dataclass
class Page:
    """Bir belgenin tek bir sayfası."""

    number: int  # 1'den başlar (insanların saydığı gibi)
    text: str


def file_hash(path: Path) -> str:
    """Dosyanın içeriğinin parmak izi.

    Dosya değişmediyse yeniden embed etmemek için kullanıyoruz — ingest'i
    ikinci kez çalıştırmak dakikalar yerine saniye sürsün diye.
    """
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


def _clean(text: str) -> str:
    """PDF'lerden gelen tipik gürültüyü temizler."""
    # Satır sonundaki tirelemeyi birleştir: "know-\nledge" -> "knowledge"
    text = re.sub(r"-\n(?=[a-zçğıöşü])", "", text)
    # Tek satır sonlarını boşluğa çevir, paragraf aralarını (çift satır) koru
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    # Üç ve fazlası boş satırı ikiye indir
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Yan yana boşlukları teke indir
    text = re.sub(r"[ \t\xa0]+", " ", text)
    return text.strip()


def load_pdf(path: Path) -> list[Page]:
    import logging

    from pypdf import PdfReader

    # pypdf, bozuk ama zararsız PDF yapılarında ekrana uyarı yağdırıyor
    # ("Ignoring wrong pointing object..."). Metin yine düzgün çıkıyor,
    # bu yüzden bu gürültüyü kapatıyoruz.
    logging.getLogger("pypdf").setLevel(logging.ERROR)

    reader = PdfReader(str(path))
    pages: list[Page] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:
            # Tek bozuk sayfa yüzünden tüm belgeyi kaybetmeyelim
            raw = ""
        cleaned = _clean(raw)
        if cleaned:
            pages.append(Page(number=i, text=cleaned))
    return pages


def load_docx(path: Path) -> list[Page]:
    """Word dosyaları sayfa kavramı taşımaz, tamamını tek 'sayfa' sayıyoruz."""
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tablolardaki metin de bilgi taşıyor (rubric'ler genelde tablo)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = _clean("\n\n".join(parts))
    return [Page(number=1, text=text)] if text else []


def load_text(path: Path) -> list[Page]:
    text = _clean(path.read_text(encoding="utf-8", errors="replace"))
    return [Page(number=1, text=text)] if text else []


_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
}


def load_document(path: Path) -> list[Page]:
    """Uzantıya göre doğru okuyucuyu seçer. Desteklenmeyen uzantı -> boş liste."""
    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        return []
    return loader(path)


def looks_like_scan(pages: list[Page]) -> bool:
    """Taranmış (görüntü) PDF'i yakalar.

    Metin katmanı olmayan bir PDF'ten neredeyse hiç karakter çıkmaz. Böyle bir
    dosyayı sessizce indekslersek bot 'bilmiyorum' der ve nedenini anlamayız.
    """
    total = sum(len(p.text) for p in pages)
    return total < 200
